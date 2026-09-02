from pathlib import Path

from docx import Document

from app.models.document import (
    BulletList,
    DocumentRequest,
    Heading,
    NumberedList,
    PageBreak,
    Paragraph,
    Table,
)
from app.renderers.base import RenderedFile

DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


class DocxRenderer:
    def render(
        self,
        request: DocumentRequest,
        output_dir: Path,
        file_name: str | None = None,
    ) -> RenderedFile:
        document = Document()
        document.core_properties.title = request.title or ""

        if request.title:
            self._add_heading(document, request.title, level=0)

        for operation in request.operations:
            if isinstance(operation, Paragraph):
                document.add_paragraph(operation.text)
            elif isinstance(operation, Heading):
                self._add_heading(document, operation.text, operation.level)
            elif isinstance(operation, BulletList):
                for item in operation.items:
                    self._add_paragraph_with_style(document, item, "List Bullet")
            elif isinstance(operation, NumberedList):
                for item in operation.items:
                    self._add_paragraph_with_style(document, item, "List Number")
            elif isinstance(operation, Table):
                table = document.add_table(rows=1, cols=len(operation.columns))
                if "Table Grid" in document.styles:
                    table.style = "Table Grid"
                for cell, value in zip(table.rows[0].cells, operation.columns, strict=True):
                    cell.text = value
                for row_values in operation.rows:
                    row = table.add_row()
                    for cell, value in zip(row.cells, row_values, strict=True):
                        cell.text = value
            elif isinstance(operation, PageBreak):
                document.add_page_break()

        file_name = file_name or request.file_name
        path = output_dir / file_name
        document.save(path)
        return RenderedFile(path=path, file_name=file_name, content_type=DOCX_CONTENT_TYPE)

    @staticmethod
    def _add_heading(document: Document, text: str, level: int) -> None:
        style_name = "Title" if level == 0 else f"Heading {level}"
        if style_name in document.styles:
            document.add_paragraph(text, style=style_name)
        else:
            document.add_paragraph(text)

    @staticmethod
    def _add_paragraph_with_style(document: Document, text: str, style_name: str) -> None:
        if style_name in document.styles:
            document.add_paragraph(text, style=style_name)
        else:
            document.add_paragraph(text)
