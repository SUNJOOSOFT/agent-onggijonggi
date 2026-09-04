import json
from pathlib import Path

from docx import Document

from app.models.document import DocumentRequest
from app.renderers.docx import DocxRenderer


def test_docx_renderer_creates_readable_document(tmp_path: Path) -> None:
    payload = json.loads((Path(__file__).parents[1] / "fixtures" / "meeting-minutes.json").read_text(encoding="utf-8"))
    rendered = DocxRenderer().render(DocumentRequest.model_validate(payload), tmp_path)

    document = Document(rendered.path)

    assert rendered.path.exists()
    assert document.core_properties.title == "AX 개발팀 회의록"
    assert "회의 개요" in [paragraph.text for paragraph in document.paragraphs]
    assert document.tables[0].cell(1, 0).text == "Document Worker"
